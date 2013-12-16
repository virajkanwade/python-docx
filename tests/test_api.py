# encoding: utf-8

"""
Test suite for the docx.api module
"""

import pytest

from mock import create_autospec, Mock, PropertyMock

from docx import parts
from docx.api import Document, _Document
from docx.opc.constants import CONTENT_TYPE as CT

from .unitutil import class_mock, var_mock


class DescribeDocument(object):

    @pytest.fixture
    def _Document_(self, request):
        return class_mock(request, 'docx.api._Document')

    @pytest.fixture
    def default_docx(self, request):
        return var_mock(request, 'docx.api._default_docx_path')

    @pytest.fixture
    def OpcPackage_(self, OpcPackage_mockery):
        return OpcPackage_mockery[0]

    @pytest.fixture
    def OpcPackage_mockery(self, request):
        OpcPackage_ = class_mock(request, 'docx.api.OpcPackage')
        pkg = OpcPackage_.open.return_value
        main_document = PropertyMock(name='main_document')
        main_document.return_value.content_type = CT.WML_DOCUMENT_MAIN
        type(pkg).main_document = main_document
        document_part = main_document.return_value
        return (OpcPackage_, pkg, main_document, document_part)

    def it_opens_a_docx_file_on_construction(self, OpcPackage_mockery,
                                             _Document_):
        # mockery ----------------------
        docx = Mock(name='docx')
        OpcPackage_, pkg, main_document, document_part = OpcPackage_mockery
        # exercise ---------------------
        doc = Document(docx)
        # verify -----------------------
        OpcPackage_.open.assert_called_once_with(docx)
        main_document.assert_called_once_with()
        _Document_.assert_called_once_with(pkg, document_part)
        assert isinstance(doc, _Document)

    def it_uses_default_if_no_file_provided(self, OpcPackage_, _Document_,
                                            default_docx):
        Document()
        OpcPackage_.open.assert_called_once_with(default_docx)

    def it_should_raise_if_not_a_docx_file(self, OpcPackage_mockery):
        # mockery ----------------------
        docx = Mock(name='docx')
        OpcPackage_, pkg, main_document, document_part = OpcPackage_mockery
        main_document.return_value.content_type = 'foobar'
        # verify -----------------------
        with pytest.raises(ValueError):
            Document(docx)


class Describe_Document(object):

    def it_provides_access_to_the_document_body(self):
        document_part = create_autospec(parts._Document)
        doc = _Document(None, document_part)
        # exercise ---------------------
        body = doc.body
        # verify -----------------------
        assert body is document_part.body

    def it_can_save_the_package(self):
        pkg, file_ = (Mock(name='pkg'), Mock(name='file_'))
        _Document(pkg, None).save(file_)
        pkg.save.assert_called_once_with(file_)
